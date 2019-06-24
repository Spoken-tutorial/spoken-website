from docx import Document
# from o import Document
from docx.enum.style import WD_STYLE_TYPE

wordDoc = Document('/home/manan/Desktop/manan.docx')
a = open("a.txt","w")
# print(wordDoc.tables)
styles = wordDoc.styles
paragraph_styles = [
    s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
]
for style in paragraph_styles:
    if style.name == 'List Bullet':
    	print ("I'm a bullet")
print(wordDoc.styles)
for table in wordDoc.tables:
    for row in table.rows:
        for cell in row.cells:
        	print(str(cell.text))
        	a.write(str(cell.text)) 
           # for paragraph_styles in cell.paragraphs:
	          #   # print (paragraph.text)
	          #   for style in paragraph_styles:
	          #   	if style.name == 'List Bullet':
	          #   		print ("I'm a bullet")
           #  print(dir(cell))
           #  print(cell.paragraphs)


